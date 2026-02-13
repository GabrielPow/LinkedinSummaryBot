import pandas as pd
#%%
from agents.agente1 import build_search
from agents.agente2 import build_summary
#%%
from docx import Document  
#%%
from agents.agente3 import build_translation
#%%
def article_building(arttype: str) -> str:
    articles = pd.read_excel("./data/articles.xlsx",header=0,usecols="B:E", dtype="object")
    index_val = articles.index[-1]
    url = articles.iloc[index_val,0]
    
    agent1_response = build_search(url)
    print(agent1_response)
    articles.loc[index_val,"Agent_1"] = agent1_response
    
    agent2_response = build_summary(agent1_response, arttype)
    print(agent2_response)
    articles.loc[index_val,"Agent_2"] = agent2_response
    
    articles.loc[index_val,"Processed"] = "Yes"
    
    with pd.ExcelWriter("./data/articles.xlsx", mode = 'a', engine = 'openpyxl', if_sheet_exists="replace") as writer:
        articles.to_excel(writer)
    
    return agent2_response
#%%
response = article_building("Ar")
#%%
document = Document()

document.add_paragraph(response)

#%%
translation = build_translation("brazillian portuguese", response)

#%%

document.add_paragraph(translation)

document.save('./output/response.docx')
#%%


